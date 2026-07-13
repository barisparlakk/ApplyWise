from __future__ import annotations

import json
import os
import re
import unicodedata
import zipfile
from io import BytesIO
from typing import Protocol
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

import pdfplumber
from docx import Document
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader

from applywise.cloudflare_ai import CloudflareAIError, CloudflareWorkersAIClient
from applywise.environment import boolean_environment

MAX_PDF_PAGES = 20
MAX_DOCX_FILES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 100_000


class ParsedResume(BaseModel):
    education: list[str] = Field(default_factory=list)
    experience: list[str] = Field(default_factory=list)
    skills: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)


class StructuredResumeProvider(Protocol):
    def extract_resume_json(self, text: str) -> str:
        pass


class ResumeExtractionError(RuntimeError):
    pass


def normalize_heading(value: str) -> str:
    normalized = unicodedata.normalize("NFKC", value).casefold().replace("&", " and ")
    return re.sub(r"[^a-z0-9+#]+", " ", normalized).strip()


SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "education": (
        "education",
        "academic background",
        "academic qualifications",
        "education and training",
        "qualifications",
    ),
    "experience": (
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "work history",
        "career history",
        "internships",
    ),
    "skills": (
        "skills",
        "technical skills",
        "technical skills and tools",
        "skills and tools",
        "technical proficiencies",
        "technologies",
        "tools and technologies",
        "competencies",
    ),
    "projects": (
        "projects",
        "selected projects",
        "personal projects",
        "academic projects",
        "project experience",
        "portfolio",
    ),
}
SECTION_NAME_LOOKUP = {
    normalize_heading(alias): section
    for section, aliases in SECTION_ALIASES.items()
    for alias in aliases
}
STOP_SECTION_HEADINGS = {
    normalize_heading(value)
    for value in (
        "summary",
        "profile",
        "about me",
        "contact",
        "certifications",
        "certificates",
        "awards",
        "publications",
        "languages",
        "interests",
        "references",
        "volunteering",
        "extracurricular activities",
    )
}


class LocalStructuredResumeProvider:
    section_names = SECTION_NAME_LOOKUP

    def extract_resume_json(self, text: str) -> str:
        sections = extract_sections_heuristically(text)
        return json.dumps(sections)


class OpenAICompatibleStructuredResumeProvider:
    def __init__(
        self,
        *,
        api_url: str,
        api_key: str,
        model: str,
        timeout_seconds: float = 30,
    ) -> None:
        self.api_url = api_url
        self.api_key = api_key
        self.model = model
        self.timeout_seconds = timeout_seconds

    def extract_resume_json(self, text: str) -> str:
        payload = {
            "model": self.model,
            "temperature": 0,
            "response_format": {"type": "json_object"},
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "Extract resume sections as JSON only. "
                        "Return exactly these keys with string-array values: "
                        "education, experience, skills, projects."
                    ),
                },
                {"role": "user", "content": text},
            ],
        }
        request = Request(
            self.api_url,
            data=json.dumps(payload).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
            },
            method="POST",
        )

        try:
            with urlopen(request, timeout=self.timeout_seconds) as response:
                response_data = json.loads(response.read())
        except (HTTPError, URLError, TimeoutError, json.JSONDecodeError) as exc:
            raise ResumeExtractionError("LLM resume extraction request failed.") from exc

        content = extract_llm_message_content(response_data)
        if content is None:
            raise ResumeExtractionError("LLM resume extraction response was not understood.")
        return content


class CloudflareStructuredResumeProvider:
    def __init__(self, client: CloudflareWorkersAIClient) -> None:
        self.client = client

    def extract_resume_json(self, text: str) -> str:
        try:
            return self.client.generate_json(
                system_prompt=(
                    "Extract the candidate's resume into the requested JSON schema. "
                    "Preserve concrete facts and the source language. Do not invent details. "
                    "Return concise string arrays for education, experience, skills, and projects."
                ),
                user_content=text[:50000],
                json_schema=ParsedResume.model_json_schema(),
                max_tokens=1800,
            )
        except CloudflareAIError as exc:
            raise ResumeExtractionError("Cloudflare resume extraction failed.") from exc


def get_structured_resume_provider() -> StructuredResumeProvider:
    provider = os.environ.get("LLM_PROVIDER", "local").strip().lower()
    if provider in {"", "local", "heuristic"}:
        return LocalStructuredResumeProvider()

    if provider == "cloudflare":
        try:
            return CloudflareStructuredResumeProvider(
                CloudflareWorkersAIClient.from_environment()
            )
        except CloudflareAIError as exc:
            raise ResumeExtractionError("Cloudflare AI is not fully configured.") from exc

    if provider in {"openai", "openai-compatible"}:
        api_url = os.environ.get("LLM_API_URL", "").strip()
        api_key = os.environ.get("LLM_API_KEY", "").strip()
        model = os.environ.get("LLM_MODEL", "").strip()
        if not api_url or not api_key or not model:
            raise ResumeExtractionError("LLM provider is not fully configured.")

        timeout_seconds = float(os.environ.get("LLM_TIMEOUT_SECONDS", "30"))
        return OpenAICompatibleStructuredResumeProvider(
            api_url=api_url,
            api_key=api_key,
            model=model,
            timeout_seconds=timeout_seconds,
        )

    raise ResumeExtractionError(f"Unsupported LLM provider: {provider}.")


def extract_llm_message_content(response_data: object) -> str | None:
    if not isinstance(response_data, dict):
        return None

    choices = response_data.get("choices")
    if isinstance(choices, list) and choices:
        first_choice = choices[0]
        if isinstance(first_choice, dict):
            message = first_choice.get("message")
            if isinstance(message, dict) and isinstance(message.get("content"), str):
                return message["content"]

    output_text = response_data.get("output_text")
    if isinstance(output_text, str):
        return output_text

    if all(key in response_data for key in ParsedResume.model_fields):
        return json.dumps(response_data)

    return None


def parse_cv_file(filename: str, content: bytes) -> str:
    try:
        lower_filename = filename.lower()
        if lower_filename.endswith(".pdf"):
            return parse_pdf(content)
        if lower_filename.endswith(".docx"):
            return parse_docx(content)
    except ResumeExtractionError:
        raise
    except Exception as exc:
        raise ResumeExtractionError("Resume file could not be parsed.") from exc
    raise ResumeExtractionError("Unsupported resume file type.")


def parse_pdf(content: bytes) -> str:
    texts: list[str] = []
    try:
        with pdfplumber.open(BytesIO(content)) as pdf:
            if len(pdf.pages) > MAX_PDF_PAGES:
                raise ResumeExtractionError(
                    f"Resume PDF must have {MAX_PDF_PAGES} pages or fewer."
                )
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
    except ResumeExtractionError:
        raise
    except Exception:
        texts = []

    if texts:
        return validate_extracted_text("\n".join(texts).strip())

    reader = PdfReader(BytesIO(content))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ResumeExtractionError(
            f"Resume PDF must have {MAX_PDF_PAGES} pages or fewer."
        )
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            texts.append(page_text)
    return validate_extracted_text("\n".join(texts).strip())


def parse_docx(content: bytes) -> str:
    validate_docx_archive(content)
    document = Document(BytesIO(content))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return validate_extracted_text("\n".join(lines))


def validate_docx_archive(content: bytes) -> None:
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            files = archive.infolist()
    except zipfile.BadZipFile as exc:
        raise ResumeExtractionError("Resume DOCX archive is invalid.") from exc
    if len(files) > MAX_DOCX_FILES:
        raise ResumeExtractionError(
            f"Resume DOCX must contain {MAX_DOCX_FILES} files or fewer."
        )
    if any(file_info.flag_bits & 0x1 for file_info in files):
        raise ResumeExtractionError("Encrypted DOCX files are not supported.")
    uncompressed_size = sum(file_info.file_size for file_info in files)
    if uncompressed_size > MAX_DOCX_UNCOMPRESSED_BYTES:
        raise ResumeExtractionError("Resume DOCX expands beyond the processing limit.")


def validate_extracted_text(text: str) -> str:
    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        raise ResumeExtractionError("Extracted resume text exceeds the processing limit.")
    return text


def extract_structured_resume(
    text: str,
    provider: StructuredResumeProvider | None = None,
) -> ParsedResume:
    extractor = provider or get_structured_resume_provider()
    last_error: Exception | None = None

    for _attempt in range(2):
        try:
            parsed = ParsedResume.model_validate_json(extractor.extract_resume_json(text))
            fallback = ParsedResume.model_validate(extract_sections_heuristically(text))
            return parsed.model_copy(
                update={
                    field: getattr(parsed, field) or getattr(fallback, field)
                    for field in ParsedResume.model_fields
                }
            )
        except (ValidationError, ValueError) as exc:
            last_error = exc
        except ResumeExtractionError:
            if provider is None and boolean_environment("AI_ALLOW_LOCAL_FALLBACK", True):
                return ParsedResume.model_validate(extract_sections_heuristically(text))
            raise

    raise ResumeExtractionError("Resume parser returned invalid structured output.") from last_error


def extract_sections_heuristically(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
    }
    current_section: str | None = None
    source_lines: list[str] = []

    for raw_line in text.splitlines():
        line = clean_resume_line(raw_line)
        if not line:
            continue
        source_lines.append(line)

        heading = split_section_heading(line)
        if heading is not None:
            current_section, inline_content = heading
            if inline_content:
                sections[current_section].append(inline_content)
            continue

        if is_non_target_heading(line):
            current_section = None
            continue

        if current_section is not None:
            sections[current_section].append(line)

    for section_name in ("education", "experience", "projects"):
        sections[section_name] = deduplicate_lines(sections[section_name])
        if not sections[section_name]:
            sections[section_name] = infer_section_lines(source_lines, section_name)

    sections["skills"] = normalize_skills(sections["skills"])
    if not sections["skills"]:
        sections["skills"] = infer_skills(source_lines)
    return sections


def clean_resume_line(value: str) -> str:
    line = unicodedata.normalize("NFKC", value).strip()
    line = re.sub(
        r"^(?:(?:[#*\s\u2022\u2023\u25e6\u2043\u2219\u2013\u2014-]+)|(?:\d{1,2}[.)]\s*))+",
        "",
        line,
    )
    return re.sub(r"\s+", " ", line).strip()


def split_section_heading(line: str) -> tuple[str, str] | None:
    normalized_line = normalize_heading(line)
    exact_section = SECTION_NAME_LOOKUP.get(normalized_line)
    if exact_section is not None:
        return exact_section, ""

    parts = re.split(r"\s*(?:[:|]|\s[-\u2013\u2014]\s)\s*", line, maxsplit=1)
    if len(parts) != 2:
        return None
    section = SECTION_NAME_LOOKUP.get(normalize_heading(parts[0]))
    if section is None:
        return None
    return section, parts[1].strip()


def is_non_target_heading(line: str) -> bool:
    return normalize_heading(line) in STOP_SECTION_HEADINGS


def deduplicate_lines(lines: list[str], *, limit: int = 30) -> list[str]:
    seen: set[str] = set()
    values: list[str] = []
    for raw_line in lines:
        line = raw_line.strip()
        key = line.casefold()
        if line and key not in seen:
            seen.add(key)
            values.append(line)
        if len(values) >= limit:
            break
    return values


EDUCATION_HINT = re.compile(
    r"\b(?:bachelor|bsc|b\.sc|master|msc|m\.sc|phd|university|college|degree|"
    r"computer engineering|data science|graduat(?:ed|ion))\b",
    re.IGNORECASE,
)
EXPERIENCE_HINT = re.compile(
    r"\b(?:intern(?:ship)?|engineer|developer|analyst|research assistant|employment|"
    r"work experience|freelance|part[- ]time)\b",
    re.IGNORECASE,
)
PROJECT_HINT = re.compile(
    r"\b(?:project|portfolio|github|repository|built|developed|implemented|designed|"
    r"prototype)\b",
    re.IGNORECASE,
)


def infer_section_lines(lines: list[str], section: str) -> list[str]:
    pattern = {
        "education": EDUCATION_HINT,
        "experience": EXPERIENCE_HINT,
        "projects": PROJECT_HINT,
    }[section]
    candidates = [
        line
        for line in lines
        if split_section_heading(line) is None
        and not is_non_target_heading(line)
        and pattern.search(line)
    ]
    return deduplicate_lines(candidates, limit=12)


KNOWN_SKILLS: tuple[tuple[str, str], ...] = (
    (r"\bpython\b", "Python"),
    (r"\bjava\b", "Java"),
    (r"\bc\+\+\b", "C++"),
    (r"\bc#\b", "C#"),
    (r"\btypescript\b", "TypeScript"),
    (r"\bjavascript\b", "JavaScript"),
    (r"\breact(?:\.js)?\b", "React"),
    (r"\bnext(?:\.js)?\b", "Next.js"),
    (r"\bnode(?:\.js)?\b", "Node.js"),
    (r"\bfastapi\b", "FastAPI"),
    (r"\bdjango\b", "Django"),
    (r"\bspring boot\b", "Spring Boot"),
    (r"\bsql\b", "SQL"),
    (r"\bpostgres(?:ql)?\b", "PostgreSQL"),
    (r"\bmysql\b", "MySQL"),
    (r"\bmongodb\b", "MongoDB"),
    (r"\bredis\b", "Redis"),
    (r"\bdocker\b", "Docker"),
    (r"\bkubernetes\b", "Kubernetes"),
    (r"\baws\b", "AWS"),
    (r"\bazure\b", "Azure"),
    (r"\bgit\b", "Git"),
    (r"\bgithub actions\b", "GitHub Actions"),
    (r"\bpandas\b", "pandas"),
    (r"\bnumpy\b", "NumPy"),
    (r"\bscikit[- ]learn\b", "scikit-learn"),
    (r"\bpytorch\b", "PyTorch"),
    (r"\btensorflow\b", "TensorFlow"),
    (r"\bmachine learning\b", "Machine Learning"),
    (r"\bdeep learning\b", "Deep Learning"),
    (r"\bcomputer vision\b", "Computer Vision"),
    (r"\bopencv\b", "OpenCV"),
    (r"\bpower bi\b", "Power BI"),
    (r"\btableau\b", "Tableau"),
    (r"\bexcel\b", "Excel"),
)


def infer_skills(lines: list[str]) -> list[str]:
    searchable_text = "\n".join(lines)
    return [
        display_name
        for pattern, display_name in KNOWN_SKILLS
        if re.search(pattern, searchable_text, re.IGNORECASE)
    ]


def normalize_skills(lines: list[str]) -> list[str]:
    seen: set[str] = set()
    skills: list[str] = []
    for line in lines:
        category_parts = line.split(":", maxsplit=1)
        skill_line = category_parts[1] if len(category_parts) == 2 else line
        for part in re.split(r"[,|;\u2022]", skill_line):
            skill = part.strip()
            key = skill.casefold()
            if skill and key not in seen:
                seen.add(key)
                skills.append(skill)
    return skills
